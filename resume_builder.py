from docx import Document
from docx.shared import Pt


def create_resume(data):

    doc = Document()

    doc.add_heading(
        data['name'],
        0
    )

    doc.add_paragraph(
        f"Email: {data['email']}"
    )

    doc.add_paragraph(
        f"Phone: {data['phone']}"
    )

    if data.get('linkedin'):
        doc.add_paragraph(
            f"LinkedIn: {data['linkedin']}"
        )

    if data.get('github'):
        doc.add_paragraph(
            f"GitHub: {data['github']}"
        )

    doc.add_heading(
        'Education',
        level=1
    )

    doc.add_paragraph(
        data['education']
    )

    doc.add_heading(
        'Skills',
        level=1
    )

    doc.add_paragraph(
        data['skills']
    )

    doc.add_heading(
        'Projects',
        level=1
    )

    doc.add_paragraph(
        data['projects']
    )

    doc.add_heading(
        'Experience',
        level=1
    )

    doc.add_paragraph(
        data['experience']
    )

    doc.add_heading(
        'Certifications',
        level=1
    )

    doc.add_paragraph(
        data['certifications']
    )

    filename = "generated_resume.docx"

    doc.save(filename)

    return filename